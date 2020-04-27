import tkinter as tk
from tkinter import filedialog
from tkinter import ttk

import numpy as np
import random
import io
import os
import openpyxl as oxl
from openpyxl.styles import Font
import csv
import PyPDF2 
import textract
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk import *
from bs4 import BeautifulSoup
import requests
import docx

from controllers import utils, log


class data_controller():
    def __init__(self, vocab, check_for_duplicates):
        self.vocab = vocab
        self.check_for_duplicates = check_for_duplicates

    def load_excel(self, excel_file="", csv_file=""):
        import_dict = []
        headings = []
        if excel_file != "":
            _file = excel_file
            wb = oxl.load_workbook(_file)
            ws = wb.active
            
            log.debug("Importing {}".format(_file))
            for i, row in enumerate(ws.rows):
                if i == 0:
                    for heading in row:
                        if heading.value != None:
                            headings.append(heading.value)
                else:
                    temp_word = {}
                    for j, heading in enumerate(headings):
                        if row[j].value != None or heading == "related_image":
                            if heading != "related_image":
                                temp_word[heading] = row[j].value
                            else:
                                image_dir = os.path.dirname(_file) + "/related_images/"
                                image_path = image_dir + str(i)+".jpg"
                                if os.path.isfile(image_path):
                                    temp_word[heading] = utils.convertToBinaryData(image_path)
                                else:
                                    temp_word[heading] = "-"
                    log.debug("Importing headings: {}".format(headings))
                    import_dict.append(temp_word)

        elif csv_file != "":
            _file = csv_file
            log.debug("Importing {}".format(_file))
            with open(csv_file, "r", encoding="utf-8-sig") as csvfile:
                rows = csv.reader(csvfile, delimiter=';')
                for i, row in enumerate(rows):
                    if i == 0:
                        for heading in row:
                            if heading != "":
                                headings.append(heading)
                    else:
                        temp_word = {}
                        for j, heading in enumerate(headings):
                            if row[j] != "" or heading == "related_image":
                                if heading != "related_image":
                                    temp_word[heading] = row[j]
                                else:
                                    image_dir = os.path.dirname(csv_file) + "/related_images/"
                                    image_path = image_dir + str(i)+".jpg"
                                    if os.path.isfile(image_path):
                                        temp_word[heading] = utils.convertToBinaryData(image_path)
                                    else:
                                        temp_word[heading] = "-"

                        import_dict.append(temp_word)

        match_count = []
        for word in import_dict:
            for heading, value in word.items():
                if heading == "transliteration" or heading == "translation":
                    duplicate_check = self.check_for_duplicates(value, heading_list=["transliteration","translation"])
            if duplicate_check != True:
                match_count.append(word)     
            else:
                pass

        if match_count != []:
            log.error("XLS/CSV Import: Found {} duplicates.".format(len(match_count)))
            message_ = '''Found {} words, which are already in vocabulary. Import anyway?
                            Press YES to import all. 
                            Press NO to import all without duplicates.
                            Press CANCEL to abort.
                            '''.format(str(len(match_count)))
            MsgBox = tk.messagebox.askyesnocancel("Found Duplicates", message_)
            if MsgBox == True:
                log.debug("Importing {} Words.".format(str(len(import_dict))))
                self.vocab.import_words_from_file(import_dict)
                return
            
            elif MsgBox == False:
                for match in match_count:
                    import_dict.remove(match)
                log.debug("Importing {} Words.".format(str(len(import_dict))))
                self.vocab.import_words_from_file(import_dict)
                return

            else:
                log.debug("Aborted Importing")
                return

        else:
            log.debug("No Duplicates Found")
            self.vocab.import_words_from_file(import_dict)

                   
    def text_extractor(self, filename, word_count=20, min_size=10, max_size=20 ):
        
        log.debug("DATA: Extracting text from {}".format(filename))
        
        if filename[-3:] == "pdf":
        
            pdfFileObj = open(filename,'rb')
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

            num_pages = pdfReader.numPages
            count = 0
            text = ""

            while count < num_pages:
                pageObj = pdfReader.getPage(count)
                count +=1
                text += pageObj.extractText()

            # Check if scanned File
            if text != "":
                text = text
            else:
                log.debug("DATA: Couldn't fetch Text, starting Textract.")
                text = textract.process(filename, method='tesseract', language='de')
        
        elif filename[-3:] == "txt":
            f = open(filename, "r", encoding="utf-8-sig")
            text = f.read()
        
        elif filename[-4:] == "docx":
            document = docx.Document(filename)
            temp_text = []
            for para in document.paragraphs:
                temp_text.append(para.text)
            text = str(temp_text)

        tokens = word_tokenize(text)
        
        punctuations = ['(',')',';',':','[',']',',']
        stop_words = stopwords.words('german') 
        keywords = [word for word in tokens if not word in stop_words and not word in punctuations]
        parametric_words = [w for w in keywords if len(w) > min_size and len(w) < max_size]

        fdist1 = FreqDist(parametric_words)
        fdist_counts = fdist1.most_common(word_count)

        final_wordlist = []

        for word in fdist_counts:
            temp_list = list(word)
            final_wordlist.append(temp_list[0])

        log.debug("DATA: Fetched {} words.".format(len(final_wordlist)))
        return(final_wordlist)


    def gen_words(self, letter_parts, word_count=30, min_size=2, max_size=6, hardness=4, foreigness=3):
        
        letters_list = np.random.randint(low = min_size, high = max_size, size = word_count)
        word = ""
        gen_words_list = []

        log.debug("DATA: Generating Word Batch. Word Count: {}, min_size: {}, max_size: {}, hardness: {}, foreigness: {}.".format(word_count, min_size, max_size, hardness, foreigness))

        for letter_num in letters_list:
                prob = random.random()*100
                word = [" "] * letter_num

                for i in range(len(word)):
                    # FIRST CHAR CHOOSER
                    if i == 0:
                        cons_prob = self.translate(hardness, 1, 10, 1, 80)
                        vowels_prob = 100-cons_prob
                        spec_prob = (vowels_prob/100)*self.translate(foreigness, 1, 10, 1, 100)
                        vowel_prob = vowels_prob-spec_prob

                        if prob <=  cons_prob:
                            word[0] = random.choice(letter_parts["consonants"])

                        elif prob > cons_prob and prob < (cons_prob+vowel_prob):
                            word[0] = random.choice(letter_parts["vowels"])

                        elif prob >= (cons_prob+vowel_prob):
                            word[0] = random.choice(letter_parts["special_vowels"])
                    
                    else:
                        prob = random.random()*100  #Generate a new value for probability

                        # Char before WAS A KONS
                        for char in letter_parts["consonants"]:
                            if char == word[i-1]:

                                cons_prob = self.translate(hardness, 1, 10, 1, 60)
                                vowels_prob = 100-cons_prob
                                spec_prob = (vowels_prob/100)*self.translate(foreigness, 1, 10, 1, 100)
                                vowel_prob = vowels_prob-spec_prob

                                if prob <= cons_prob:
                                    word[i] = random.choice(letter_parts["consonants"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["consonants"])
                                    except:
                                        word[i] = random.choice(letter_parts["consonants"])

                                elif prob > cons_prob and prob < (cons_prob+vowel_prob):
                                    word[i] = random.choice(letter_parts["vowels"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["vowels"])
                                    except:
                                        word[i] = random.choice(letter_parts["vowels"])

                                elif prob >= (cons_prob+vowel_prob):
                                    word[i] = random.choice(letter_parts["special_vowels"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["special_vowels"])
                                    except:
                                        word[i] = random.choice(letter_parts["special_vowels"])

                            else: 
                                pass
                        
                        # Char before  WAS A VOWEL
                        for char in letter_parts["vowels"]:
                            if char == word[i-1]:

                                cons_prob = self.translate(hardness, 1, 10, 1, 85)
                                vowels_prob = 100-cons_prob
                                spec_prob = (vowels_prob/100)*self.translate(foreigness, 1, 10, 1, 100)
                                vowel_prob = vowels_prob-spec_prob

                                if prob <= cons_prob:
                                    word[i] = random.choice(letter_parts["consonants"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["consonants"])
                                    except:
                                        word[i] = random.choice(letter_parts["consonants"])

                                elif prob > cons_prob and prob < (cons_prob+vowel_prob):
                                    word[i] = random.choice(letter_parts["vowels"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["vowels"])
                                    except:
                                        word[i] = random.choice(letter_parts["vowels"])

                                elif prob >= (cons_prob+vowel_prob):
                                    word[i] = random.choice(letter_parts["special_vowels"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["special_vowels"])
                                    except:
                                        word[i] = random.choice(letter_parts["special_vowels"])
                            else:
                                pass


                        # Char before WAS A SPECIAL VOWEL
                        for char in letter_parts["special_vowels"]:
                            if char == word[i-1]:

                                cons_prob = self.translate(hardness, 1, 10, 1, 85)
                                vowels_prob = 100-cons_prob
                                spec_prob = (vowels_prob/100)*self.translate(foreigness, 1, 10, 1, 100)
                                vowel_prob = vowels_prob-spec_prob

                                if prob <= cons_prob:
                                    word[i] = random.choice(letter_parts["consonants"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["consonants"])
                                    except:
                                        word[i] = random.choice(letter_parts["consonants"])

                                elif prob > cons_prob and prob < (cons_prob+vowel_prob):
                                    word[i] = random.choice(letter_parts["vowels"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["vowels"])
                                    except:
                                        word[i] = random.choice(letter_parts["vowels"])

                                elif prob >= (cons_prob+vowel_prob):
                                    word[i] = random.choice(letter_parts["special_vowels"])
                                    try:
                                        while word[i] == word[i-1] and word[i] == word[i-2]:
                                            word[i] = random.choice(letter_parts["special_vowels"])
                                    except:
                                        word[i] = random.choice(letter_parts["special_vowels"])
                            else:
                                pass

                        

                word = "".join(word)
                gen_words_list.append(word)
                word = []

        gen_words_set = list(set(gen_words_list))
        return gen_words_set
    

    def translate(self, value, leftMin, leftMax, rightMin, rightMax):
        # Figure out how 'wide' each range is
        leftSpan = leftMax - leftMin
        rightSpan = rightMax - rightMin

        # Convert the left range into a 0-1 range (float)
        valueScaled = float(value - leftMin) / float(leftSpan)

        # Convert the 0-1 range into a value in the right range.
        return rightMin + (valueScaled * rightSpan)

    
    def get_language_from_web(self, url="https://www.1000mostcommonwords.com/"):
        
        headers = requests.utils.default_headers()
        headers.update({ 'User-Agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:52.0) Gecko/20100101 Firefox/52.0'})
        
        try:
            req = requests.get(url, headers)
        except:
            log.error("DATA: Requesting {} failed.".format(url))

        soup = BeautifulSoup(req.content, 'html.parser')
        links = {}
        log.debug("DATA: Scraping data from {}".format(url))

        if url == "https://www.1000mostcommonwords.com/":

            output = soup.find_all("a", attrs={"style" : "color: #0000ff;"})

            for link in output:
                links[link.find(text=True)] = link.get("href")
            
            return(links)
    

    def get_words_from_web(self, url, start_count=0, end_count=999):
        headers = requests.utils.default_headers()
        headers.update({ 'User-Agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:52.0) Gecko/20100101 Firefox/52.0'})

        try:
            req = requests.get(url, headers)
        except:
            log.error("DATA: Requesting {} failed.".format(url))

        soup = BeautifulSoup(req.content, 'html.parser')

        output = soup.find_all("tr")

        if "1000mostcommonwords" in url:

            words_list = []

            for row in output:
                words_dict = {}
                cell = row.find_all("td")
                words_dict["translation"]=cell[1].text 
                words_dict["english"]=cell[2].text

                words_list.append(words_dict)
        log.debug("DATA: Importing {} words.".format(len(words_list[int(start_count):int(end_count)])))
        return(words_list[int(start_count):int(end_count)])
    

    def export_vocabulary_as_file(self, filename, vocabulary, formats, columns):
        for format_ in formats:
            if format_ == "CSV":
                
                if ".csv" not in filename:
                    output_name = str(filename) + ".csv"
                else:
                    output_name = filename

                with open(output_name, 'w', newline='') as csvfile:
                    writer = csv.writer(csvfile, delimiter=';')
                    writer.writerow(columns)
                    for word_ in vocabulary: 
                        temp_attr = []
                        for column in columns:
                            temp_attr.append(word_[column])  
                        writer.writerow(temp_attr)

            elif format_ == "XLSX":
                if ".xlsx" not in filename:
                    output_name = str(filename) + ".xlsx"
                else:
                    output_name = filename

                wb = oxl.Workbook()
                ws = wb.active
                ws.title = "vocabulary"
                ft = Font(bold=True)
                
                ws.append(columns)
                for word_ in vocabulary: 
                        temp_attr = []
                        for i, column in enumerate(columns): 
                            if column != "related_image":
                                temp_attr.append(word_[column])  
                        ws.append(temp_attr)

                for cell in ws["1"]:
                    cell.font = ft

                wb.save(filename = output_name)

            elif format_ == "TXT":
                if ".txt" not in filename:
                    output_name = str(filename) + ".txt"
                else:
                    output_name = filename
            
            log.debug("DATA: Exporting {}.".format(output_name))
    
    
    def export_batch(self, generated_word_list, filename):
        log.debug("DATA: Exporting {}".format(filename))
        if filename[-3:] == "csv":
            try:
                with open(filename, mode='w', newline='') as csv_file:
                    csv_writer = csv.writer(csv_file, delimiter=";")
                    for word_ in generated_word_list:
                        csv_writer.writerow([word_])
            except:
                log.error("DATA: Exporting {} failed.".format(filename))
        
        elif filename[-4:] == "xlsx":
            wb = oxl.Workbook() 
            sheet = wb.active 

            for word_ in generated_word_list:
                sheet.append([word_])
            try:
                wb.save(filename) 
            except:
                log.error("DATA: Exporting {} failed.".format(filename))

        elif filename[-4:] == "docx":
            document = docx.Document()

            document.add_heading('Vocabulary', 0)
            for word_ in generated_word_list:
                document.add_paragraph(word_)
            try:
                document.save(filename)
            except:
                log.error("DATA: Exporting {} failed.".format(filename))

        elif filename[-3:] == "txt":
            try:
                txt_file = open(filename, "w")
                for word_ in generated_word_list:
                    txt_file.write(word_ + "\n") 
                txt_file.close() 
            except:
                log.error("DATA: Exporting {} failed.".format(filename))
